"""
Generate response.docx — formal response letter addressing all reviewer comments.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPATH = os.path.join(os.path.dirname(__file__), 'response.docx')

FONT_NAME = 'Times New Roman'

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.03)
    section.bottom_margin = Cm(2.03)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)


# ── Helpers ──────────────────────────────────────────────────────────────────

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, italic=False, indent=False, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = space_after
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_comment_response(comment_text, response_text):
    """Add a reviewer comment (italic, indented) followed by a response."""
    # Comment
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    label = p.add_run("Comment: ")
    label.font.name = FONT_NAME
    label.font.size = Pt(10)
    label.bold = True
    label.italic = True
    body = p.add_run(comment_text)
    body.font.name = FONT_NAME
    body.font.size = Pt(10)
    body.italic = True

    # Response
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(12)
    label2 = p2.add_run("Response: ")
    label2.font.name = FONT_NAME
    label2.font.size = Pt(10)
    label2.bold = True
    body2 = p2.add_run(response_text)
    body2.font.name = FONT_NAME
    body2.font.size = Pt(10)


# ═════════════════════════════════════════════════════════════════════════════
# CONTENT
# ═════════════════════════════════════════════════════════════════════════════

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Response to Reviewer Comments")
run.font.name = FONT_NAME
run.font.size = Pt(16)
run.bold = True

# Manuscript title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "Context Without Contact: Top-Down Information Flow from Shared "
    "Modulation in Uncoupled Prebiotic Systems"
)
run.font.name = FONT_NAME
run.font.size = Pt(11)
run.italic = True

# Journal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run("Manuscript submitted to Royal Society Open Science")
run.font.name = FONT_NAME
run.font.size = Pt(10)

# ── Preamble ─────────────────────────────────────────────────────────────────

add_para(
    "We thank the reviewers for their constructive and insightful comments. "
    "We have carefully addressed every point raised. Below, we respond to "
    "each comment individually and describe the specific changes made. "
    "All modifications to the manuscript are highlighted in red font in the "
    "revised manuscript and supplementary information."
)

# ═════════════════════════════════════════════════════════════════════════════
# REVIEWER 1
# ═════════════════════════════════════════════════════════════════════════════

add_heading("Reviewer 1", level=1)

# ── Major 1 ──────────────────────────────────────────────────────────────────

add_heading("Major Comment 1: Strengthen the Introduction", level=2)

add_comment_response(
    "The article is structured into three sections: the Introduction provides "
    "background, the Results section outlines the methodology and numerical "
    "experiments, and the Discussion interprets the findings in the context of "
    "the origin of life. I recommend a more thorough treatment of the underlying "
    "themes. In particular, the Introduction would benefit from additional detail "
    "and references to prior work that has considered top-down causation in the "
    "context of the origin of life and evolutionary transitions. Greater attention "
    "should also be given to the role of logistic map models in computational "
    "biology and the application of information-theoretic measures. These additions "
    "would substantially strengthen the Introduction, which in its current form is "
    "insufficient.",

    "We have substantially expanded the Introduction with three new paragraphs:\n\n"
    "(a) Top-down causation in origin of life and evolutionary transitions: "
    "We added a paragraph reviewing prior work by Walker et al., Ellis, "
    "Szathm\u00e1ry, Flack, and others on top-down causation in the context of "
    "major evolutionary transitions and the origin of life. This paragraph "
    "now provides the conceptual context for our study and identifies the gap "
    "it fills (references 4\u20139 in the revised manuscript).\n\n"
    "(b) Logistic maps in computational biology: We added a paragraph covering "
    "the use of logistic map models in population dynamics and computational "
    "biology, including their role as minimal models exhibiting deterministic "
    "chaos. We also acknowledged that synchronization of chaotic systems under "
    "common external forcing is a well-characterized phenomenon in nonlinear "
    "dynamics, citing Boccaletti et al. (2002) (references 14, 31\u201336).\n\n"
    "(c) Information-theoretic measures: We added a paragraph reviewing the "
    "application of transfer entropy and related information-theoretic measures "
    "to biological systems, multi-scale dynamics, and causality detection, "
    "including discussion of strengths and limitations of TE as a causality "
    "proxy (references 17\u201324)."
)

# ── Major 2 ──────────────────────────────────────────────────────────────────

add_heading("Major Comment 2: Add a dedicated Methodology section", level=2)

add_comment_response(
    "I strongly recommend adding a dedicated methodology section that clearly "
    "presents the model and parameter choices, some of which appear arbitrary "
    "(e.g., r_min, r_max, K\u2080, A, \u03c9). This would help demonstrate the "
    "robustness of the argument.",

    "We have created a new \u201cModel and Methods\u201d section, placed between the "
    "Introduction and Results. This section now contains:\n\n"
    "(a) Model description: The full model definition, including equations, "
    "variable definitions, and the ensemble mean field definition, has been "
    "moved from the former Results section into this new section. The Results "
    "section now opens directly with findings.\n\n"
    "(b) Parameter justification table: We added a summary table (Table 1) "
    "listing every model parameter (r_min, r_max, K\u2080, A, \u03c9, N, T, "
    "number of bins, embedding lag k) together with its value and a brief "
    "rationale grounded in biological motivation, literature conventions, or "
    "sensitivity analysis.\n\n"
    "(c) TE estimation procedure: We documented the histogram-based transfer "
    "entropy estimation method in full detail, including bin count selection, "
    "lag selection rationale (referencing Supplementary Figure S1), the circular "
    "shift surrogate testing procedure, number of surrogates, and significance "
    "threshold. A reader can now reproduce the TE computation from the Methods "
    "section alone."
)

# ── Major 3 ──────────────────────────────────────────────────────────────────

add_heading("Major Comment 3: Disentangle top-down causation from "
            "common-driver forcing", level=2)

add_comment_response(
    "From a methodological standpoint, I am not convinced that the Author "
    "demonstrates convincingly that the alignment of individual units with "
    "the mean field is due to top-down causation rather than to forcing caused "
    "by the environmental modulation A. [...] I encourage the Author to "
    "supplement their analysis by isolating constraint-driven order from "
    "mean-field-driven order. Conditional transfer entropy could be useful for "
    "this purpose. Moreover, I encourage the Author to pay particular attention "
    "to the value of A/K\u2080, which determines the extent to which the quadratic "
    "term influences the forcing from the environmental constraint. Combined "
    "with a deeper analysis of Figure 1 \u2014 for instance, why TE increases "
    "monotonically with A, and why the confidence interval rises sharply when "
    "A \u2248 50 \u2014 this would substantially strengthen the Author\u2019s argument.",

    "We thank the reviewer for this important suggestion, which has "
    "substantially strengthened the manuscript. We have addressed all three "
    "sub-points:\n\n"
    "(a) Conditional transfer entropy (CTE) analysis: We computed CTE in both "
    "directions (M \u2192 x and x \u2192 M), conditioning on the environmental signal "
    "K\u2099, across the same parameter sweep as Figure 1B (A = 0 to 60). Results "
    "are presented in a new Figure 3 alongside the unconditional TE for direct "
    "comparison. The analysis reveals a two-component structure: the "
    "environmental signal accounts for a large share of the unconditional TE "
    "(as expected, since it is the shared driver), but a genuine non-zero "
    "residual CTE(M\u2192x|K) persists across all modulation amplitudes. This "
    "residual peaks at approximately 0.36 bits at intermediate A values, "
    "demonstrating that the nonlinear aggregation of heterogeneous chaotic "
    "units generates emergent statistical structure in the mean field that "
    "carries predictive information about individual trajectories beyond what "
    "the environment alone provides. The two components \u2014 environmental "
    "scaffolding and genuine emergent information flow \u2014 are discussed "
    "explicitly in both the Results and Discussion sections.\n\n"
    "(b) A/K\u2080 ratio sweep: We performed a parameter sweep over A/K\u2080 values "
    "with multiple K\u2080 settings (50, 100, 200). The results (new Figure 4) "
    "show that TE collapses onto a single universal curve when plotted against "
    "A/K\u2080, confirming this ratio as the natural control parameter governing "
    "the transition from the nonlinear (quadratic-dominated) regime to the "
    "near-linear regime. This is discussed in the Results section.\n\n"
    "(c) Explanation of Figure 1B features: We added explicit explanations for "
    "the two features the reviewer highlighted. First, TE(M\u2192x) increases "
    "monotonically with A because stronger modulation forces individual "
    "trajectories to track the environment more closely, making the mean field "
    "an increasingly accurate predictor of each unit\u2019s next state. Second, "
    "the sharp rise in the confidence interval near A \u2248 50 corresponds to the "
    "regime boundary (A/K\u2080 \u2248 0.5) where the system transitions from nonlinear "
    "dynamics to near-complete environmental entrainment; in this transition "
    "zone, small differences in initial conditions and growth rates lead to "
    "large variability in how quickly individual realizations lock onto the "
    "modulated attractor."
)

# ── Minor 3 ──────────────────────────────────────────────────────────────────

add_heading("Minor Comment 3: Unify terminology", level=2)

add_comment_response(
    'Please unify the terminology: "ensemble mean" (p. 4) vs. "mean field" '
    '(p. 5), "population mean" (p. 6), and so on.',

    'We have standardized the terminology to "mean field" throughout the '
    "entire manuscript (main text, figure captions, and supplementary "
    "materials). All instances of \u201censemble mean\u201d and \u201cpopulation mean\u201d "
    "have been replaced."
)

# ── Minor 4 ──────────────────────────────────────────────────────────────────

add_heading("Minor Comment 4: Add omega inset to Figure 1B", level=2)

add_comment_response(
    "In Figure 1B, I suggest adding an inset illustrating how TE varies "
    "as a function of \u03c9.",

    "We have added an inset to Figure 1B showing TE(M\u2192x) as a function "
    "of the modulation frequency \u03c9. The sweep covers values spanning an "
    "order of magnitude with other parameters held at defaults. A brief "
    "description has been added to the Results text."
)

# ── Minor 5 ──────────────────────────────────────────────────────────────────

add_heading("Minor Comment 5: Clarify the saturation / soft-bound statement",
            level=2)

add_comment_response(
    'The Author writes that "The saturation was not due to numerical clipping; '
    'although a soft bound..." Please explain.',

    "We have expanded this passage to clarify three points: (a) what the soft "
    "bound of (0, 2K\u2080) is and why it was imposed (to prevent biologically "
    "unrealistic negative populations or extreme overshoots while preserving "
    "the natural dynamics), (b) how often the bound was approached at different "
    "A values (fewer than 10% of time steps at A = 60, and negligibly at lower "
    "A), and (c) why this rules out clipping as the driver of TE saturation "
    "(since the bound is rarely active, the saturation in TE reflects the "
    "genuine information-theoretic ceiling of environmental entrainment rather "
    "than an artifact)."
)

# ── Minor 6 ──────────────────────────────────────────────────────────────────

add_heading("Minor Comment 6: Write out SNR and R\u00b2 formulas", level=2)

add_comment_response(
    "Please explicitly write out the SNR and R\u00b2 formulas actually computed.",

    "We have added explicit formulas for both metrics in the Methods section. "
    "R\u00b2 is defined as the standard coefficient of determination between "
    "M\u2099 and x\u2081,\u2099. SNR is defined as Var(M\u2099) / Var(\u03b5), "
    "where \u03b5 denotes the residuals from the linear regression of x\u2081,\u2099 "
    "on M\u2099. Both formulas include full variable definitions."
)

# ── Minor 7 ──────────────────────────────────────────────────────────────────

add_heading("Minor Comment 7: Fix the GitHub repo URL", level=2)

add_comment_response(
    "The address of the GitHub repo is incorrect.",

    "The GitHub repository URL has been corrected in the revised manuscript."
)

# ── Minor 8 ──────────────────────────────────────────────────────────────────

add_heading("Minor Comment 8: Fix Supplementary Figure S3 caption", level=2)

add_comment_response(
    'Supplementary Information, p. 4: The caption should indicate '
    '"Return maps of [...] A = {0, 20, 40, 60}".',

    "The caption for Supplementary Figure S3 has been corrected to read "
    "A = {0, 20, 40, 60}, matching the actual figure panels."
)

# ═════════════════════════════════════════════════════════════════════════════
# REVIEWER 2
# ═════════════════════════════════════════════════════════════════════════════

add_heading("Reviewer 2", level=1)

add_comment_response(
    "The starting point of this work is the fact that a set of chaotic systems "
    "can synchronize even in the absence of mutual coupling when driven by a "
    "common external signal [...] This is a well-known phenomenon, characterized "
    "in the classical literature on chaotic systems. [...] I think this work "
    "should be mentioned, in particular for its section 5: Boccaletti, S., "
    "Kurths, J., Osipov, G., Valladares, D. L., & Zhou, C. S. (2002). The "
    "synchronization of chaotic systems. Physics Reports, 366(1\u20132), 1\u2013101. "
    "It goes without saying that it would be appreciated, at least in the "
    "continuation of this research line, to show the implications of the "
    "concept of synchronization in a simulation including some concrete aspect "
    "of biochemistry or biophysics.",

    "We thank the reviewer for this important reference. We have made two "
    "additions to the manuscript:\n\n"
    "(a) We now explicitly acknowledge in the Introduction that the "
    "synchronization of chaotic systems under common external forcing is a "
    "well-characterized phenomenon in nonlinear dynamics, citing Boccaletti "
    "et al. (2002) as reference 36. The relevant sentence reads: \u201cMore broadly, "
    "the synchronization of chaotic systems under common external forcing is a "
    "well-characterized phenomenon in nonlinear dynamics (36); even without "
    "mutual coupling, a shared driving signal can entrain individually chaotic "
    "units into coherent collective behavior.\u201d\n\n"
    "(b) We have added a future-direction statement at the end of the Discussion "
    "noting that a natural extension of this work would be to test these "
    "principles in models incorporating concrete aspects of prebiotic "
    "biochemistry or biophysics \u2014 for example, populations of protocells with "
    "explicit metabolic kinetics under cyclic environmental forcing."
)

# ═════════════════════════════════════════════════════════════════════════════
# REVIEWER 3
# ═════════════════════════════════════════════════════════════════════════════

add_heading("Reviewer 3", level=1)

add_comment_response(
    "One minor suggestion would be to sharpen, perhaps in the Discussion, the "
    "distinction between externally imposed boundary conditions and internally "
    'generated coordination, to further preempt potential misunderstandings '
    'regarding "top-down" causation.',

    "We have added a new paragraph in the Discussion that explicitly "
    "distinguishes two modes of coordination often conflated under the umbrella "
    'of "top-down causation": (i) externally imposed boundary conditions \u2014 '
    "such as the modulated carrying capacity K\u2099 in our model \u2014 which constrain "
    "all units from outside the system, and (ii) internally generated "
    "coordination arising from coupling among units themselves (e.g., "
    "autocatalytic feedback, cross-catalysis, or metabolic networks). The "
    "paragraph explains that our model demonstrates how external boundary "
    "conditions alone can produce the informational signatures of top-down "
    "influence, and connects this distinction to the CTE analysis: the "
    "environmental component reflects the boundary-condition mechanism, while "
    "the genuine residual CTE reflects emergent statistical structure from "
    "nonlinear aggregation \u2014 a form of collective information processing that "
    "goes beyond passive entrainment. This clarifies that the \u2018top-down\u2019 signal "
    "in our model is not a claim of autonomous macroscopic agency, but rather "
    "a demonstration that externally structured environments can generate the "
    "informational preconditions from which internally coordinated systems may "
    "subsequently emerge."
)

# ═════════════════════════════════════════════════════════════════════════════

doc.save(OUTPATH)
print(f"Response letter saved to: {OUTPATH}")
