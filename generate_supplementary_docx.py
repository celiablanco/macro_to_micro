"""
Generate the revised supplementary information as a Word document (.docx)
with changes highlighted in red.
Mirrors generate_supplementary_pdf.py but outputs .docx for Google Docs import.
"""
import os
from html.parser import HTMLParser
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIGDIR = os.path.join(os.path.dirname(__file__), 'figures')
OUTPATH = os.path.join(os.path.dirname(__file__), 'revised_supplementary.docx')

# ─────────────────────────────────────────────────────────────────────────────
# HTML-to-runs parser (same as paper generator)
# ─────────────────────────────────────────────────────────────────────────────

class RunSpec:
    def __init__(self, text, bold=False, italic=False, subscript=False,
                 superscript=False, red=False):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.subscript = subscript
        self.superscript = superscript
        self.red = red

class MarkupParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.runs = []
        self.tag_stack = []
        self.current_text = ""

    def _flush(self):
        if self.current_text:
            bold = any(t == 'b' for t in self.tag_stack)
            italic = any(t == 'i' for t in self.tag_stack)
            sub = any(t == 'sub' for t in self.tag_stack)
            sup = any(t == 'sup' for t in self.tag_stack)
            red = any(t == 'font_red' for t in self.tag_stack)
            self.runs.append(RunSpec(self.current_text, bold, italic, sub, sup, red))
            self.current_text = ""

    def handle_starttag(self, tag, attrs):
        self._flush()
        if tag == 'font':
            color = dict(attrs).get('color', '')
            self.tag_stack.append('font_red' if color == 'red' else 'font_other')
        else:
            self.tag_stack.append(tag)

    def handle_endtag(self, tag):
        self._flush()
        target = 'font_red' if tag == 'font' else tag
        target2 = 'font_other' if tag == 'font' else None
        for i in range(len(self.tag_stack) - 1, -1, -1):
            if self.tag_stack[i] == target or (target2 and self.tag_stack[i] == target2):
                self.tag_stack.pop(i)
                break

    def handle_data(self, data):
        self.current_text += data

    def get_runs(self):
        self._flush()
        return self.runs


def parse_markup(text):
    parser = MarkupParser()
    parser.feed(text)
    return parser.get_runs()


def R(text):
    return f'<font color="red">{text}</font>'

def B(text):
    return f'<b>{text}</b>'


# ─────────────────────────────────────────────────────────────────────────────
# Document helpers
# ─────────────────────────────────────────────────────────────────────────────

RED = RGBColor(0xFF, 0x00, 0x00)
FONT_NAME = 'Times New Roman'

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.03)
    section.bottom_margin = Cm(2.03)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(10)


def _apply_runs(paragraph, markup_text, font_size=Pt(10)):
    runs = parse_markup(markup_text)
    for rs in runs:
        run = paragraph.add_run(rs.text)
        run.font.size = font_size
        run.font.name = FONT_NAME
        if rs.bold:
            run.bold = True
        if rs.italic:
            run.italic = True
        if rs.subscript:
            run.font.subscript = True
        if rs.superscript:
            run.font.superscript = True
        if rs.red:
            run.font.color.rgb = RED


def add_para(text, font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    _apply_runs(p, text, font_size)
    return p


def add_caption(text, font_size=Pt(9)):
    return add_para(text, font_size=font_size, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_after=Pt(6))


def add_figure(filename, width_inches=4.0):
    path = os.path.join(FIGDIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc.paragraphs[-1]


# ═════════════════════════════════════════════════════════════════════════════
# CONTENT
# ═════════════════════════════════════════════════════════════════════════════

# Title
add_para("Supplementary Information:", font_size=Pt(14),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

add_para(
    "Context Without Contact: Top-Down Information Flow from Shared "
    "Modulation in Uncoupled Prebiotic Systems",
    font_size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=Pt(18))

add_para(B("This PDF file includes:"), space_after=Pt(3))
add_para("Supporting Figures S1 to S3", space_after=Pt(24))

# ═════════════════════════════════════════════════════════════════════════════
# FIGURE S1
# ═════════════════════════════════════════════════════════════════════════════
add_figure('figS1.png', 4.0)
add_caption(
    B("Figure S1. Mean transfer entropy (TE) as a function of embedding lag k for modulation "
      "amplitude A=25. ") +
    "Top-down TE (M\u2192x, circles) measures information from the macroscopic "
    + R("mean field") + " to an individual unit, while bottom-up TE (x\u2192M, squares) measures "
    "the reverse direction. Values are averaged over 10 random seeds. Parameters: N=1000, "
    "K\u2080=100, r<sub>i</sub>\u2208[3.9,4.0], T=10,000 steps (first 1,000 discarded), "
    "15-bin histograms.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# FIGURE S2
# ═════════════════════════════════════════════════════════════════════════════
add_figure('figS2.png', 6.0)
add_caption(
    B("Figure S2. Directional information flow from macroscopic to microscopic dynamics under "
      "different modulation profiles. ") +
    "Transfer entropy (TE) between the " + R("mean field") + " M<sub>n</sub> and a representative "
    "unit x<sub>n</sub>, computed in both directions: M \u2192 x (blue) and x \u2192 M (gray), "
    "for a population of N = 1000 uncoupled logistic units evolving under carrying capacity "
    "modulated by (A,C) logistic increase, and (B,D) inverse logistic. In all cases "
    "K<sub>0</sub> = 100, A = 25, r<sub>i</sub> \u2208 [3.9, 4.0], and steepness is k = 0.05, "
    "simulated over 10 independent simulations, each run for T = 10,000 steps (first 1,000 "
    "discarded). Top row: The black curve shows the external modulation K<sub>n</sub>, the blue "
    "line corresponds to the " + R("mean field") + " M, sampled every 200 steps, and the shaded "
    "gray band indicates the range of individual states at each sampled point. Bottom row: TE "
    "estimated using a histogram-based method with lag k = 2 and 15 bins. Shaded regions "
    "indicate 95% confidence intervals.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# FIGURE S3 (MINOR-8: corrected caption A = {0, 20, 40, 60})
# ═════════════════════════════════════════════════════════════════════════════
add_figure('figS3.png', 5.0)
add_caption(
    B("Figure S3. Structural modulation induces low-dimensional macroscopic dynamics. ") +
    "Return maps of (M<sub>n</sub>, M<sub>n+1</sub>) for "
    + R("A = {0, 20, 40, 60}") +
    ", N = 1,000, K<sub>0</sub> = 100, r<sub>i</sub> \u2208 [3.9, 4.0], and T = 10,000 steps "
    "(first 1,000 discarded). The dashed line shows the identity M<sub>n+1</sub> = M<sub>n</sub> "
    "and the gray curve shows the single-unit logistic map for comparison. At A=0, points cluster "
    "near (K<sub>0</sub>/2, K<sub>0</sub>/2). As A increases, shared modulation widens the "
    "distribution of M<sub>n</sub>, producing a diagonal spread.")

# Build document
doc.save(OUTPATH)
print(f"Revised supplementary saved to: {OUTPATH}")
